#!/usr/bin/env python
# -*- coding: UTF-8 -*-
from docx import Document
from docx import shared
from pathlib import Path

word_file_path = './src'

doc = Document()


def add_picture(picture):
    doc.add_picture(picture, width=shared.Inches(1))


add_picture('kail.png')
doc.save(Path(word_file_path, 'kail.docx'))
