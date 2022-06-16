#!/usr/bin/env python
# -*- coding: UTF-8 -*-
from docx import Document
from docx.shared import RGBColor
from pathlib import Path

word_file_path = './src'

doc = Document()
my_context = '''
你叉叉，穷哈哈，
哈利波特骑着扫帚飞，
而我的扫帚却只能清理垃圾堆，
什么是，快乐星期？穷哈哈
'''


def add_context(context):
    para = doc.add_paragraph().add_run(context)

    # 设置字体格式
    para.font.name = '仿宋'
    # 设置下划线
    para.font.underline = True
    # 设置颜色
    para.font.color.rgb = RGBColor(255, 128, 128)
    # 其他设置参考官方文档
    # https://python-docx.readthedocs.io/en/latest/api/text.html#run-objects


add_context(my_context)
doc.save(Path(word_file_path, 'new2.docx'))
